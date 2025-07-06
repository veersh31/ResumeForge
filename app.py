import os
import base64
import io
from dash import Dash, html, dcc, Input, Output, State, callback, no_update, ALL
import dash_bootstrap_components as dbc
from docx import Document
import PyPDF2
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
import spacy
import re
from docx.shared import Pt, Inches
import time
import json
from collections import Counter
import random


nltk.download('punkt')
nltk.download('averaged_perceptron_tagger')
nltk.download('maxent_ne_chunker')
nltk.download('words')
nltk.download('stopwords')
nlp = spacy.load("en_core_web_sm")

app = Dash(__name__, external_stylesheets=[
    dbc.themes.BOOTSTRAP,
    "https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0/css/all.min.css",
    "https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;500;600;700&display=swap"
])


app.index_string = '''
<!DOCTYPE html>
<html>
    <head>
        {%metas%}
        <title>{%title%}</title>
        {%favicon%}
        {%css%}
        <style>
            body {
                font-family: 'Poppins', sans-serif;
                background-color: #f8f9fa;
            }
            .fade-in {
                animation: fadeIn 0.5s ease-in;
            }
            @keyframes fadeIn {
                from { opacity: 0; }
                to { opacity: 1; }
            }
            .card {
                transition: transform 0.3s ease;
            }
            .card:hover {
                transform: translateY(-5px);
            }
            .upload-area {
                transition: all 0.3s ease;
            }
            .upload-area:hover {
                border-color: #0d6efd !important;
                background-color: #e9ecef !important;
            }
            .template-card {
                cursor: pointer;
                transition: all 0.3s ease;
            }
            .template-card:hover {
                transform: scale(1.03);
                box-shadow: 0 10px 20px rgba(0,0,0,0.1);
            }
            .template-card.selected {
                border: 2px solid #0d6efd;
            }
            .skill-badge {
                margin-right: 5px;
                margin-bottom: 5px;
                display: inline-block;
            }
            .progress-bar {
                transition: width 0.5s ease;
            }
        </style>
    </head>
    <body>
        {%app_entry%}
        <footer>
            {%config%}
            {%scripts%}
            {%renderer%}
        </footer>
    </body>
</html>
'''

ACTION_VERBS = [
    "achieved", "developed", "led", "managed", "created", "implemented",
    "designed", "launched", "increased", "decreased", "improved", "optimized",
    "streamlined", "coordinated", "facilitated", "negotiated", "resolved",
    "analyzed", "evaluated", "researched", "planned", "executed", "delivered"
]

POWER_WORDS = [
    "expert", "innovative", "strategic", "proactive", "dynamic", "versatile",
    "resourceful", "efficient", "effective", "skilled", "experienced",
    "knowledgeable", "professional", "dedicated", "motivated", "driven"
]

RESUME_TEMPLATES = [
    {
        "id": "modern",
        "name": "Modern",
        "description": "Clean and contemporary design with a focus on readability",
        "icon": "fas fa-briefcase"
    },
    {
        "id": "professional",
        "name": "Professional",
        "description": "Traditional format with a professional appearance",
        "icon": "fas fa-user-tie"
    },
    {
        "id": "creative",
        "name": "Creative",
        "description": "Stand out with a unique and creative layout",
        "icon": "fas fa-palette"
    },
    {
        "id": "minimal",
        "name": "Minimal",
        "description": "Simple and elegant with minimal distractions",
        "icon": "fas fa-minus-circle"
    }
]


INDUSTRY_SKILLS = {
    "technology": [
        "Python", "JavaScript", "Java", "C++", "SQL", "React", "Node.js", 
        "AWS", "Docker", "Kubernetes", "Machine Learning", "Data Analysis",
        "DevOps", "CI/CD", "Git", "RESTful APIs", "Microservices"
    ],
    "finance": [
        "Financial Analysis", "Budgeting", "Forecasting", "Risk Management",
        "Investment Banking", "Portfolio Management", "Financial Modeling",
        "Excel", "Bloomberg Terminal", "Financial Reporting", "Tax Planning"
    ],
    "healthcare": [
        "Patient Care", "Medical Records", "HIPAA Compliance", "Clinical Research",
        "Healthcare Administration", "Medical Billing", "Healthcare Informatics",
        "Public Health", "EHR Systems", "Medical Terminology"
    ],
    "marketing": [
        "Digital Marketing", "SEO", "Content Strategy", "Social Media Marketing",
        "Email Marketing", "Analytics", "Brand Management", "Market Research",
        "Campaign Management", "CRM", "Adobe Creative Suite"
    ],
    "education": [
        "Curriculum Development", "Classroom Management", "Student Assessment",
        "Educational Technology", "Special Education", "Lesson Planning",
        "Student Engagement", "Educational Leadership", "Differentiated Instruction"
    ]
}

def parse_contents(contents):
    content_type, content_string = contents.split(',')
    decoded = base64.b64decode(content_string)
    return decoded

def extract_text_from_pdf(pdf_bytes):
    pdf_file = io.BytesIO(pdf_bytes)
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(docx_bytes):
    doc = Document(io.BytesIO(docx_bytes))
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def create_docx_from_text(text, filename, template="modern"):
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    if template == "modern":
        style.font.name = 'Arial'
        style.font.size = Pt(10)
        title = doc.add_paragraph()
        title.alignment = 1
        title_run = title.add_run("PROFESSIONAL RESUME")
        title_run.bold = True
        title_run.font.size = Pt(16)
        doc.add_paragraph()
        
    elif template == "professional":
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)
        header = doc.add_paragraph()
        header.alignment = 1
        header_run = header.add_run("PROFESSIONAL RESUME")
        header_run.bold = True
        header_run.font.size = Pt(14)
        doc.add_paragraph()
        
    elif template == "creative":
        style.font.name = 'Verdana'
        style.font.size = Pt(10)
        header = doc.add_paragraph()
        header.alignment = 1
        header_run = header.add_run("RESUME")
        header_run.bold = True
        header_run.font.size = Pt(18)
        doc.add_paragraph()
        
    elif template == "minimal":
        style.font.name = 'Helvetica'
        style.font.size = Pt(10)
        header = doc.add_paragraph()
        header.alignment = 0
        header_run = header.add_run("RESUME")
        header_run.bold = True
        header_run.font.size = Pt(14)
        doc.add_paragraph()
    
    sections = text.split('\n\n')
    for section in sections:
        if section.strip():
            if section.isupper() or ':' in section:
                p = doc.add_paragraph()
                p.add_run(section.strip()).bold = True
                p.style = 'Heading 1'
            else:
                p = doc.add_paragraph()
                p.add_run(section.strip())
            doc.add_paragraph()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    doc_byte = io.BytesIO()
    doc.save(doc_byte)
    doc_byte.seek(0)
    return doc_byte

def improve_clarity(text):
    sentences = sent_tokenize(text)
    improved_sentences = []
    changes = []
    
    for sentence in sentences:
        original = sentence
        sentence = re.sub(r'\b(in order to)\b', 'to', sentence)
        sentence = re.sub(r'\b(due to the fact that)\b', 'because', sentence)
        sentence = re.sub(r'\b(at this point in time)\b', 'now', sentence)
        sentence = re.sub(r'\b(very|really|quite|rather)\s+', '', sentence)
        
        if original != sentence:
            changes.append(f"Improved clarity: '{original}' → '{sentence}'")
        
        improved_sentences.append(sentence)
    
    return ' '.join(improved_sentences), changes

def add_action_verbs(text):
    sentences = sent_tokenize(text)
    improved_sentences = []
    changes = []
    
    for sentence in sentences:
        original = sentence
        doc = nlp(sentence)
        has_action_verb = any(token.pos_ == "VERB" for token in doc)
        
        if not has_action_verb and len(sentence.split()) > 3:
            if any(word in sentence.lower() for word in ["project", "initiative", "program"]):
                sentence = f"Led {sentence}"
            elif any(word in sentence.lower() for word in ["analysis", "research", "study"]):
                sentence = f"Conducted {sentence}"
            elif any(word in sentence.lower() for word in ["improvement", "increase", "growth"]):
                sentence = f"Achieved {sentence}"
            
            if original != sentence:
                changes.append(f"Added action verb: '{original}' → '{sentence}'")
        
        improved_sentences.append(sentence)
    
    return ' '.join(improved_sentences), changes

def optimize_for_ats(text):
    common_keywords = [
        "project management", "team leadership", "communication",
        "problem solving", "analytical skills", "collaboration",
        "strategic planning", "data analysis", "innovation"
    ]
    
    changes = []
    for keyword in common_keywords:
        if keyword.lower() not in text.lower():
            text += f"\n• {keyword.title()}"
            changes.append(f"Added ATS keyword: '{keyword.title()}'")
    
    return text, changes

def enhance_professional_tone(text):
    replacements = {
        "got": "obtained",
        "made": "created",
        "did": "performed",
        "helped": "assisted",
        "worked on": "contributed to",
        "looked at": "analyzed",
        "fixed": "resolved",
        "started": "initiated",
        "ended": "completed"
    }
    
    changes = []
    for informal, formal in replacements.items():
        if re.search(r'\b' + informal + r'\b', text, flags=re.IGNORECASE):
            text = re.sub(r'\b' + informal + r'\b', formal, text, flags=re.IGNORECASE)
            changes.append(f"Enhanced tone: '{informal}' → '{formal}'")
    
    return text, changes

def add_quantifiable_achievements(text):
    sentences = sent_tokenize(text)
    improved_sentences = []
    changes = []
    
    for sentence in sentences:
        original = sentence
        if any(word in sentence.lower() for word in ["increased", "decreased", "improved", "reduced"]):
            if not any(char.isdigit() for char in sentence):
                if "increased" in sentence.lower():
                    sentence = sentence.replace("increased", "increased by 25%")
                elif "decreased" in sentence.lower():
                    sentence = sentence.replace("decreased", "decreased by 20%")
                elif "improved" in sentence.lower():
                    sentence = sentence.replace("improved", "improved by 30%")
                elif "reduced" in sentence.lower():
                    sentence = sentence.replace("reduced", "reduced by 15%")
                
                if original != sentence:
                    changes.append(f"Added quantifiable metric: '{original}' → '{sentence}'")
        
        improved_sentences.append(sentence)
    
    return ' '.join(improved_sentences), changes

def extract_skills(text):
    words = word_tokenize(text.lower())
    stop_words = set(nltk.corpus.stopwords.words('english'))
    filtered_words = [word for word in words if word.isalnum() and word not in stop_words]
    word_freq = Counter(filtered_words)
    
    all_skills = []
    for industry_skills in INDUSTRY_SKILLS.values():
        all_skills.extend([skill.lower() for skill in industry_skills])
    
    found_skills = []
    for skill in all_skills:
        if skill in text.lower():
            found_skills.append(skill.title())
    
    if not found_skills:
        skill_indicators = ["proficient in", "experienced with", "knowledge of", "skills in", "expertise in"]
        for indicator in skill_indicators:
            if indicator in text.lower():
                pattern = f"{indicator}\s+([a-zA-Z\s]+)"
                matches = re.findall(pattern, text.lower())
                if matches:
                    found_skills.append(matches[0].strip().title())
    
    if not found_skills:
        industry_keywords = {
            "technology": ["software", "programming", "code", "developer", "engineer", "tech"],
            "finance": ["finance", "banking", "accounting", "financial", "investment"],
            "healthcare": ["health", "medical", "patient", "clinical", "healthcare"],
            "marketing": ["marketing", "advertising", "brand", "campaign", "social media"],
            "education": ["education", "teaching", "student", "classroom", "curriculum"]
        }
        
        for industry, keywords in industry_keywords.items():
            if any(keyword in text.lower() for keyword in keywords):
                suggested_skills = INDUSTRY_SKILLS[industry][:5]
                found_skills.extend(suggested_skills)
                break
        
        if not found_skills:
            found_skills = ["Communication", "Problem Solving", "Teamwork", "Time Management", "Leadership"]
    
    return list(set(found_skills))

def analyze_resume(text):
    analysis = {
        "strengths": [],
        "improvements": [],
        "suggestions": []
    }
    
    doc = nlp(text)
    action_verbs_count = sum(1 for token in doc if token.pos_ == "VERB")
    if action_verbs_count > 10:
        analysis["strengths"].append(f"Good use of action verbs ({action_verbs_count} found)")
    else:
        analysis["improvements"].append(f"Consider adding more action verbs (only {action_verbs_count} found)")
    
    if re.search(r'\d+%', text) or re.search(r'\$\d+', text):
        analysis["strengths"].append("Contains quantifiable achievements")
    else:
        analysis["improvements"].append("Add more quantifiable achievements with numbers and percentages")
    
    informal_words = ["got", "made", "did", "helped", "worked on", "looked at", "fixed", "started", "ended"]
    informal_count = sum(1 for word in informal_words if re.search(r'\b' + word + r'\b', text, flags=re.IGNORECASE))
    if informal_count == 0:
        analysis["strengths"].append("Professional tone throughout")
    else:
        analysis["improvements"].append(f"Replace {informal_count} informal words with more professional alternatives")
    
    ats_keywords = ["project management", "team leadership", "communication", "problem solving", 
                   "analytical skills", "collaboration", "strategic planning", "data analysis", "innovation"]
    found_keywords = [keyword for keyword in ats_keywords if keyword.lower() in text.lower()]
    if len(found_keywords) >= 5:
        analysis["strengths"].append(f"Good ATS optimization with {len(found_keywords)} keywords")
    else:
        analysis["improvements"].append(f"Add more industry keywords for ATS optimization (only {len(found_keywords)} found)")
    
    long_sentences = [sent for sent in sent_tokenize(text) if len(sent.split()) > 20]
    if len(long_sentences) > 5:
        analysis["improvements"].append(f"Simplify {len(long_sentences)} long sentences for better readability")
    else:
        analysis["strengths"].append("Good sentence length and clarity")
    
    skills = extract_skills(text)
    if skills:
        analysis["strengths"].append(f"Identified {len(skills)} relevant skills")
    
    if len(text.split()) < 300:
        analysis["suggestions"].append("Consider expanding your resume with more details about your experience")
    if not re.search(r'@', text):
        analysis["suggestions"].append("Add your email address to the resume")
    if not re.search(r'\d{3}[-.]?\d{3}[-.]?\d{4}', text):
        analysis["suggestions"].append("Add your phone number to the resume")
    
    return analysis

def extract_job_keywords(job_description):
    if not job_description or job_description.strip() == "":
        return []
    
    doc = nlp(job_description.lower())
    keywords = []
    
    for chunk in doc.noun_chunks:
        if len(chunk.text.split()) <= 3:
            keywords.append(chunk.text.strip())
    
    for ent in doc.ents:
        if ent.label_ in ["ORG", "PRODUCT", "GPE", "PERSON"]:
            keywords.append(ent.text.strip())
    
    for token in doc:
        if token.pos_ == "VERB" and not token.is_stop:
            keywords.append(token.lemma_)
    
    skill_indicators = ["proficient in", "experience with", "knowledge of", "familiarity with", 
                        "expertise in", "ability to", "skills in", "required", "preferred"]
    
    for indicator in skill_indicators:
        if indicator in job_description.lower():
            pattern = f"{indicator}\s+([a-zA-Z\s,]+)(?:\s|$)"
            matches = re.findall(pattern, job_description.lower())
            for match in matches:
                skills = [skill.strip() for skill in match.split(',')]
                keywords.extend(skills)
    
    return list(set(keywords))

def optimize_for_job(text, job_description):
    if not job_description or job_description.strip() == "":
        return text, []
    
    job_keywords = extract_job_keywords(job_description)
    if not job_keywords:
        return text, []
    
    changes = []
    missing_keywords = []
    for keyword in job_keywords:
        if keyword.lower() not in text.lower():
            missing_keywords.append(keyword)
    
    if missing_keywords:
        if "SKILLS" not in text.upper() and "SKILLS:" not in text.upper():
            text += "\n\nSKILLS:\n"
            changes.append("Added a SKILLS section to highlight job-relevant skills")
        
        skills_section = "\n".join([f"• {keyword}" for keyword in missing_keywords[:5]])
        text += f"\n{skills_section}"
        changes.append(f"Added {len(missing_keywords[:5])} job-relevant keywords to the skills section")
    
    job_title_match = re.search(r'(job title|position|role):\s*([^\n]+)', job_description.lower())
    if job_title_match:
        job_title = job_title_match.group(2).strip()
        if job_title not in text.lower():
            if "PROFESSIONAL SUMMARY" in text.upper() or "OBJECTIVE" in text.upper():
                sentences = sent_tokenize(text)
                for i, sentence in enumerate(sentences):
                    if "PROFESSIONAL SUMMARY" in sentence.upper() or "OBJECTIVE" in sentence.upper():
                        if i+1 < len(sentences):
                            next_sentence = sentences[i+1]
                            if job_title not in next_sentence.lower():
                                sentences[i+1] = f"Seeking a {job_title} position where I can " + next_sentence.lower()
                                changes.append(f"Added job title '{job_title}' to professional summary")
                                break
                text = " ".join(sentences)
            else:
                text = f"PROFESSIONAL SUMMARY\n\nSeeking a {job_title} position where I can utilize my skills and experience.\n\n" + text
                changes.append(f"Added a professional summary targeting the {job_title} position")
    
    return text, changes


app.layout = dbc.Container([
    
    dbc.Row([
        dbc.Col([
            html.H1("Resume Forge", className="text-center my-4 text-primary"),
            html.P("Enhance your resume with AI-powered improvements", className="text-center text-muted mb-4 h5"),
        ], width=12)
    ], className="bg-light py-4 rounded shadow-sm"),
    
    
    dbc.Row([
        
        dbc.Col([
            dbc.Card([
                dbc.CardHeader("Upload Your Resume", className="bg-primary text-white"),
                dbc.CardBody([
                    dcc.Upload(
                        id='upload-resume',
                        children=html.Div([
                            html.I(className="fas fa-cloud-upload-alt fa-3x mb-3"),
                            html.Div('Drag and Drop or '),
                            html.A('Select Files', className="text-primary")
                        ], className="text-center"),
                        style={
                            'width': '100%',
                            'height': '200px',
                            'lineHeight': '200px',
                            'borderWidth': '2px',
                            'borderStyle': 'dashed',
                            'borderRadius': '10px',
                            'textAlign': 'center',
                            'margin': '10px',
                            'backgroundColor': '#f8f9fa',
                            'cursor': 'pointer'
                        },
                        className="upload-area",
                        multiple=False
                    ),
                    html.Div(id='output-upload', className="mt-3"),
                ])
            ], className="shadow-sm"),
        ], width=6, className="mb-4"),
        
        
        dbc.Col([
            dbc.Card([
                dbc.CardHeader("Enhancement Options", className="bg-primary text-white"),
                dbc.CardBody([
                    dbc.Checklist(
                        id='enhancement-options',
                        options=[
                            {"label": "Improve clarity and conciseness", "value": "clarity"},
                            {"label": "Add action verbs and power words", "value": "action_verbs"},
                            {"label": "Optimize for ATS (Applicant Tracking Systems)", "value": "ats"},
                            {"label": "Enhance professional tone", "value": "tone"},
                            {"label": "Add quantifiable achievements", "value": "achievements"},
                            {"label": "Target for specific job", "value": "job_target"},
                        ],
                        value=[],
                        className="mb-4",
                    ),
                    html.Div([
                        dbc.Button(
                            "Analyze Resume",
                            id="enhance-button",
                            color="primary",
                            className="w-100",
                            size="lg"
                        ),
                    ], className="d-grid gap-2"),
                ])
            ], className="shadow-sm"),
        ], width=6, className="mb-4"),
    ]),
    
    
    html.Div([
        dbc.Row([
            dbc.Col([
                dbc.Card([
                    dbc.CardHeader("Job Description", className="bg-info text-white"),
                    dbc.CardBody([
                        dbc.Textarea(
                            id="job-description-input",
                            placeholder="Paste the job description here to target your resume for this position...",
                            style={"height": "200px"},
                            className="mb-3"
                        ),
                        html.Div([
                            dbc.Button(
                                "Use Job Description",
                                id="use-job-description-button",
                                color="info",
                                className="w-100",
                            ),
                        ], className="d-grid gap-2"),
                    ])
                ], className="shadow-sm"),
            ], width=12, className="mb-4"),
        ]),
    ], id="job-description-section", style={"display": "none"}),
    
    
    dbc.Row([
        dbc.Col([
            dbc.Card([
                dbc.CardHeader("Choose a Resume Template", className="bg-primary text-white"),
                dbc.CardBody([
                    dbc.Row([
                        dbc.Col([
                            dbc.Card([
                                dbc.CardBody([
                                    html.Div([
                                        html.I(className=f"{template['icon']} fa-2x mb-3 text-primary"),
                                        html.H5(template["name"], className="card-title"),
                                        html.P(template["description"], className="card-text small text-muted"),
                                    ], className="text-center")
                                ])
                            ], className="template-card", id={"type": "template-card", "index": i})
                        ], width=3) for i, template in enumerate(RESUME_TEMPLATES)
                    ]),
                    html.Div(id="selected-template-display", className="mt-3 text-center"),
                    dcc.RadioItems(
                        id="template-selector",
                        options=[{"label": template["name"], "value": template["id"]} for template in RESUME_TEMPLATES],
                        value="modern",
                        inline=True,
                        className="mt-3"
                    ),
                ])
            ], className="shadow-sm mb-4"),
        ], width=12)
    ]),
    
    
    dbc.Row([
        dbc.Col([
            html.Div(id="progress-container", className="d-none"),
            dbc.Progress(id="progress-bar", value=0, className="mb-3", style={"height": "10px"}),
            html.Div(id="progress-status", className="text-center text-muted mb-4")
        ], width=12)
    ]),
    
    # Results Section
    dbc.Row([
        dbc.Col([
            html.Div(id="enhancement-output", className="mt-4"),
        ], width=12)
    ]),
    
    dbc.Row([
        dbc.Col([
            html.Div(id="analysis-output", className="mt-4"),
        ], width=12)
    ]),
    
    # Skills Section
    dbc.Row([
        dbc.Col([
            html.Div(id="skills-output", className="mt-4"),
        ], width=12)
    ]),
    
    # Download Section
    dbc.Row([
        dbc.Col([
            html.Div(id="download-section", className="mt-4 text-center"),
        ], width=12)
    ]),
    
    # Store for the enhanced text
    dcc.Store(id='enhanced-text-store'),
    dcc.Store(id='extracted-skills-store'),
    
    # Footer
    dbc.Row([
        dbc.Col([
            html.Hr(),
            html.P("Resume Forge - AI-Powered Resume Enhancement", className="text-center text-muted"),
            html.P("No API keys required - All processing done locally", className="text-center text-muted small"),
        ], width=12)
    ], className="mt-4"),
], fluid=True, className="py-4")

@callback(
    Output('output-upload', 'children'),
    Input('upload-resume', 'contents'),
    State('upload-resume', 'filename')
)
def update_output(contents, filename):
    if contents is None:
        return html.Div("No file uploaded yet", className="text-muted")
    
    return html.Div([
        dbc.Alert([
            html.I(className="fas fa-check-circle me-2"),
            f"File uploaded: {filename}"
        ], color="success", className="mb-0"),
        html.P("Your resume is ready for analysis!", className="text-muted mt-2 mb-0")
    ])

@callback(
    Output("progress-container", "className"),
    Output("progress-bar", "value"),
    Output("progress-status", "children"),
    Output('enhancement-output', 'children'),
    Output('analysis-output', 'children'),
    Output('skills-output', 'children'),
    Output('download-section', 'children'),
    Output('enhanced-text-store', 'data'),
    Output('extracted-skills-store', 'data'),
    Input('enhance-button', 'n_clicks'),
    State('upload-resume', 'contents'),
    State('upload-resume', 'filename'),
    State('enhancement-options', 'value'),
    State('template-selector', 'value'),
    State('job-description-input', 'value')
)
def enhance_resume(n_clicks, contents, filename, enhancement_options, selected_template, job_description):
    if n_clicks is None or contents is None:
        return "d-none", 0, "", None, None, None, None, None, None
    
    
    progress_container = ""
    
    
    file_bytes = parse_contents(contents)
    
    
    if filename.endswith('.pdf'):
        resume_text = extract_text_from_pdf(file_bytes)
    elif filename.endswith('.docx'):
        resume_text = extract_text_from_docx(file_bytes)
    else:
        return "d-none", 0, "", html.Div("Unsupported file format"), None, None, None, None, None
    
    
    progress = 20
    progress_status = "Extracting text from resume..."
    
    
    skills = extract_skills(resume_text)
    
    
    progress = 40
    progress_status = "Analyzing resume content..."
    
    
    enhanced_text = resume_text
    all_changes = []
    
    for option in enhancement_options:
        if option == "clarity":
            enhanced_text, changes = improve_clarity(enhanced_text)
            all_changes.extend(changes)
        elif option == "action_verbs":
            enhanced_text, changes = add_action_verbs(enhanced_text)
            all_changes.extend(changes)
        elif option == "ats":
            enhanced_text, changes = optimize_for_ats(enhanced_text)
            all_changes.extend(changes)
        elif option == "tone":
            enhanced_text, changes = enhance_professional_tone(enhanced_text)
            all_changes.extend(changes)
        elif option == "achievements":
            enhanced_text, changes = add_quantifiable_achievements(enhanced_text)
            all_changes.extend(changes)
        elif option == "job_target" and job_description:
            enhanced_text, changes = optimize_for_job(enhanced_text, job_description)
            all_changes.extend(changes)
    
    
    progress = 60
    progress_status = "Enhancing resume content..."
    
    
    analysis = analyze_resume(enhanced_text)
    
    
    if "job_target" in enhancement_options and job_description:
        job_keywords = extract_job_keywords(job_description)
        if job_keywords:
            
            found_keywords = [keyword for keyword in job_keywords if keyword.lower() in enhanced_text.lower()]
            match_percentage = len(found_keywords) / len(job_keywords) * 100
            
            if match_percentage >= 70:
                analysis["strengths"].append(f"Strong match with job requirements ({match_percentage:.0f}% of keywords found)")
            elif match_percentage >= 40:
                analysis["strengths"].append(f"Moderate match with job requirements ({match_percentage:.0f}% of keywords found)")
            else:
                analysis["improvements"].append(f"Low match with job requirements (only {match_percentage:.0f}% of keywords found)")
            
            
            missing_keywords = [keyword for keyword in job_keywords if keyword.lower() not in enhanced_text.lower()]
            if missing_keywords:
                analysis["suggestions"].append(f"Consider adding these job-specific keywords: {', '.join(missing_keywords[:5])}")
    
    
    progress = 80
    progress_status = "Preparing results..."
    
    
    enhancement_output = html.Div([
        dbc.Card([
            dbc.CardHeader("Enhanced Resume", className="bg-success text-white"),
            dbc.CardBody([
                html.Pre(enhanced_text, style={
                    'whiteSpace': 'pre-wrap',
                    'fontFamily': 'Arial, sans-serif',
                    'fontSize': '14px',
                    'lineHeight': '1.6',
                    'padding': '20px',
                    'backgroundColor': '#f8f9fa',
                    'borderRadius': '5px',
                    'border': '1px solid #dee2e6'
                })
            ])
        ], className="shadow-sm")
    ])
    
    
    analysis_output = html.Div([
        dbc.Card([
            dbc.CardHeader("Resume Analysis", className="bg-primary text-white"),
            dbc.CardBody([
                
                html.Div([
                    html.H5("Strengths", className="text-success mb-3"),
                    html.Ul([html.Li(item, className="mb-2") for item in analysis["strengths"]], 
                           className="list-unstyled") if analysis["strengths"] 
                    else html.P("No specific strengths identified.", className="text-muted")
                ], className="mb-4"),
                
                
                html.Div([
                    html.H5("Areas for Improvement", className="text-danger mb-3"),
                    html.Ul([html.Li(item, className="mb-2") for item in analysis["improvements"]], 
                           className="list-unstyled") if analysis["improvements"] 
                    else html.P("No specific improvements needed.", className="text-muted")
                ], className="mb-4"),
                
                
                html.Div([
                    html.H5("Additional Suggestions", className="text-primary mb-3"),
                    html.Ul([html.Li(item, className="mb-2") for item in analysis["suggestions"]], 
                           className="list-unstyled") if analysis["suggestions"] 
                    else html.P("No additional suggestions.", className="text-muted")
                ], className="mb-4"),
                
                html.Hr(),
                
                
                html.Div([
                    html.H5("Changes Made", className="mb-3"),
                    html.Ul([html.Li(change, className="mb-2") for change in all_changes], 
                           className="list-unstyled") if all_changes 
                    else html.P("No specific changes were made to the text.", className="text-muted")
                ])
            ])
        ], className="shadow-sm")
    ])
    
    
    skills_output = html.Div([
        dbc.Card([
            dbc.CardHeader("Identified Skills", className="bg-info text-white"),
            dbc.CardBody([
                html.Div([
                    dbc.Badge(skill, color="info", className="skill-badge p-2 m-1") 
                    for skill in skills
                ], className="d-flex flex-wrap"),
                html.Hr(),
                html.H5("Industry Skills", className="mb-3"),
                html.Div([
                    dbc.Accordion([
                        dbc.AccordionItem([
                            html.Div([
                                dbc.Badge(skill, color="secondary", className="skill-badge p-2 m-1") 
                                for skill in industry_skills
                            ], className="d-flex flex-wrap")
                        ], title=industry.capitalize())
                        for industry, industry_skills in INDUSTRY_SKILLS.items()
                    ], start_collapsed=True)
                ])
            ])
        ], className="shadow-sm")
    ])
    

    download_button = html.Div([
        dbc.Card([
            dbc.CardHeader("Download Enhanced Resume", className="bg-success text-white"),
            dbc.CardBody([
                html.P("Your resume has been enhanced and is ready to download.", className="mb-3"),
                dbc.Button(
                    "Download Resume",
                    id="download-button",
                    color="success",
                    className="me-2",
                    size="lg"
                ),
                dcc.Download(id="download-resume")
            ])
        ], className="shadow-sm")
    ])
    
    
    progress = 100
    progress_status = "Analysis complete!"
    
    return progress_container, progress, progress_status, enhancement_output, analysis_output, skills_output, download_button, enhanced_text, skills

@callback(
    Output("selected-template-display", "children"),
    Input("template-selector", "value")
)
def update_selected_template(selected_template):
    template = next((t for t in RESUME_TEMPLATES if t["id"] == selected_template), None)
    if template:
        return html.Div([
            html.I(className=f"{template['icon']} me-2"),
            f"Selected: {template['name']} Template"
        ], className="text-primary")
    return ""

@callback(
    Output({"type": "template-card", "index": ALL}, "className"),
    Input("template-selector", "value")
)
def update_template_cards(selected_template):
    return [
        "template-card" + (" selected" if template["id"] == selected_template else "")
        for template in RESUME_TEMPLATES
    ]

@callback(
    Output("download-resume", "data"),
    Input("download-button", "n_clicks"),
    State("enhanced-text-store", "data"),
    State("template-selector", "value"),
    prevent_initial_call=True
)
def download_resume(n_clicks, enhanced_text, selected_template):
    if enhanced_text is None:
        return None
    
    
    doc_byte = create_docx_from_text(enhanced_text, "enhanced_resume.docx", selected_template)
    
    return dcc.send_bytes(
        doc_byte.getvalue(),
        "enhanced_resume.docx",
        base64=True
    )

@callback(
    Output("job-description-section", "style"),
    Input("enhancement-options", "value"),
    prevent_initial_call=True
)
def toggle_job_description_section(enhancement_options):
    if "job_target" in enhancement_options:
        return {"display": "block"}
    return {"display": "none"}

if __name__ == '__main__':
    app.run_server(debug=True) 
